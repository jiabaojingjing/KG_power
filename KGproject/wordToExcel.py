from docx import Document
import win32com.client as wc
from openpyxl import Workbook
import os
import win32file
import shutil
path = r'C:\Users\86136\Desktop\电力缺陷\通辽调研资料\变电五项管理规定的206册细则Word版\变电评价管理规定细则'
docxpath = r'C:\Users\86136\Desktop\电力缺陷\通辽调研资料\变电五项管理规定的206册细则Word版\变电评价管理规定细则docx'
excelpath = r'C:\Users\86136\Desktop\电力缺陷\通辽调研资料\变电五项管理规定的206册细则Word版\变电评价管理规定细则excel'

def docxToexcel():
    docxfileset=get_all(docxpath);
    print(docxfileset)
    #excelfileset=get_all(excelpath);
    for h in docxfileset:
        pathdocx=docxpath+"\\"+h
        pathexcel = excelpath+"\\"+os.path.splitext(h)[0]+".xlsx"
        wb = Workbook()
        sheet = wb.active
        document = Document(pathdocx)
        tables = document.tables
        #print(tables)
        global str1,excelset
        excelset=[]
        for j in range(len(tables)):
            print( " tableid: "+str(j))
            for i in range(len(tables[j].rows)):
                excelset.clear()
                print(" row: "+str(len(tables[j].rows)) +" rowid: "+str(i))
                for k in range(len(tables[j].columns)):
                    print("colnum: "+str(len(tables[j].columns)) + "  colid  " + str(k))
                    try:
                        str1 = tables[j].cell(i, k).text
                        if str1 not in excelset:
                            excelset.append(str1)
                    except:
                        print("出错啦")
                        wb.close()
                if len(excelset)>0:
                    print(excelset)
                    sheet.append(excelset)
        wb.save(pathexcel)
        wb.close()
def test():

    document = Document(r'C:\Users\86136\Desktop\电力缺陷\通辽调研资料\变电五项管理规定的206册细则Word版\变电评价管理规定细则docx\国家电网公司变电评价管理规定（试行） 第32分册 组合电器检修策略.docx')
    tables2 = document.tables
    str1 = tables2[0].cell(208, 10).text
    print(str1)
    print(str(len(tables2[0].rows)))

def docToDocx():
    docfileset=get_all(path)
    docxfileset = get_all(docxpath)
    for i in docfileset:
        filedoc = path + "\\" + i
        if os.path.splitext(i)[1]==".doc":
            if i not in docxfileset:
                filedocx=docxpath+"\\"+os.path.splitext(i)[0]+".docx"
                word = wc.Dispatch("Word.Application")
                print(filedoc)
                doc = word.Documents.Open(filedoc)
                # 上面的地方只能使用完整绝对地址，相对地址找不到文件，且，只能用“\\”，不能用“/”，哪怕加了 r 也不行，涉及到将反斜杠看成转义字符。
                doc.SaveAs(filedocx, 12, False, "", True, "", False, False, False, False)  # 转换后的文件,12代表转换后为docx文件
                # doc.SaveAs(r"F:\\***\\***\\appendDoc\\***.docx", 12)#或直接简写
                # 注意SaveAs会打开保存后的文件，有时可能看不到，但后台一定是打开的
                doc.Close
                word.Quit
        elif os.path.splitext(i)[1]==".docx":
            shutil.copy(filedoc, docxpath)

        print(os.path.splitext(i)[0])


def get_all(cwd):
    global row
    global col
    global folderflag
    global subfolder
    fileset=[]
    get_dir = os.listdir(cwd)
    for i in get_dir:
        #print(i)
        sub_dir = os.path.join(cwd,i)
        #print(sub_dir)
        if os.path.isdir(sub_dir):
            subfolder=True
            row += 1
            get_all(sub_dir)
        else:
            if win32file.GetFileAttributesW(sub_dir)==32:
                filename=os.path.splitext(i)[0]
                fileset.append(i)
    return fileset

if __name__ == "__main__":
    #get_all(r'C:\Users\86136\Desktop\电力缺陷\通辽调研资料\变电五项管理规定的206册细则Word版\变电评价管理规定细则')
    # wordtoexcel()
    #docToDocx()
    docxToexcel()
    #test()