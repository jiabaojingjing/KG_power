
import time
import csv
import docx
from os import walk
import xlrd
import traceback
import os
import time
from datetime import datetime
from elasticsearch import Elasticsearch
from elasticsearch.helpers import bulk
result = []
row=0
col=0


superfoldernode=foldernode=subfoldernode=super_folder=folder=sub_folder=filename=""
document=firsttitle=secondtitle=thirdtitle=fourthtitle=documentnode=firstnode=secondnode=thirdnode =fournode=""
#workbook = xlwt.Workbook(encoding='utf-8')
#worksheet = workbook.add_sheet('My Sheet')
path=r"C:\Users\86136\Desktop\文档内容整理\文档"




class ElasticObj:
    def __init__(self, index_name,index_type,ip ="127.0.0.1"):
        '''

        :param index_name: 索引名称
        :param index_type: 索引类型
        '''
        self.index_name =index_name
        self.index_type = index_type
        # 无用户名密码状态
        #self.es = Elasticsearch([ip])
        #用户名密码状态
        self.es = Elasticsearch([ip],http_auth=('elastic', 'password'),port=9200)

    # def create_index(self, index_name="ott", index_type="ott_type"):
    #     '''
    #     创建索引,创建索引名称为ott，类型为ott_type的索引
    #     :param ex: Elasticsearch对象
    #     :return:
    #     '''
    #     # 创建映射
    #     _index_mappings = {
    #
    #     }
    #     if self.es.indices.exists(index=self.index_name) is not True:
    #         res = self.es.indices.create(index=self.index_name, body=_index_mappings)
    #         # print(res)

    def wipe_line_break(self,str):
        return str.replace("\n", "").replace(" ", "")

    def delete_BS(self,str):
        return str.replace(" ", "")

    def IndexData(self):
        es = Elasticsearch()
        csvdir = 'D:/work/ElasticSearch/exportExcels'
        filenamelist = []
        for (dirpath, dirnames, filenames) in walk(csvdir):
            filenamelist.extend(filenames)
            break
        total = 0
        for file in filenamelist:
            csvfile = csvdir + '/' + file
            self.Index_Data_FromCSV(csvfile,es)
            total += 1
            print (total)
            time.sleep(10)

    def Index_Data_FromCSV(self,csvfile):
        '''
        从CSV文件中读取数据，并存储到es中
        :param csvfile: csv文件，包括完整路径
        :return:
        '''
        list = csv.ReadCSV(csvfile)
        index = 0
        doc = {}
        for item in list:
            if index > 1:#第一行是标题
                doc['title'] = item[0]
                doc['link'] = item[1]
                doc['date'] = item[2]
                doc['source'] = item[3]
                doc['keyword'] = item[4]
                res = self.es.index(index=self.index_name, doc_type=self.index_type, body=doc)
                print(res['created'])
            index += 1
            # print index

    def Index_Data(self):
        '''
        数据存储到es
        :return:
        '''
        list = [
            {   "date": "2020-09-13",
                "source": "慧聪网",
                "link": "http://info.broadcast.hc360.com/2017/09/130859749974.shtml",
                "keyword": "电视",
                "title": "长期急救周期性负载",
                "name": "变压器长时间在环境温度较高，或者超过额定电流条件下运行。这种运行方式将不同程度缩短变压器的寿命，应尽量减少这种运行方式出现的机会；必须采用时，应尽量缩短超过额定电流运行时间，降低超过额定电流的倍数，投入备用冷却器。长期急救周期性负载状态下的负载电流、温度限值及最长时间见表2。在长期急救周期性负载运行期间，应有负载电流记录，并计算该运行期间的平均相对老化率。"
             },
            {   "date": "2020-09-15",
                "source": "中国文明网",
                "link": "http://www.wenming.cn/xj_pd/yw/201709/t20170913_4421323.shtml",
                "keyword": "电视",
                "title": "短期急救负载",
                "name": "变压器短时间大幅度超过额定电流条件下运行，这种负载可能导致绕组热点温度达到危险的程度，使绝缘强度暂时下降，应投入（包括备用冷却器在内的）全部冷却器（制造厂另有规定的除外），并尽量压缩负载，减少时间，一般不超过0.5h。短期急救负载状态下的负载电流、温度限值及最长时间见表2。在短期急救负载运行期间，应有详细的负载电流记录，并计算该运行期间的相对老化率。"
             }
              ]
        for item in list:
            print(self.index_name)
            res = self.es.index(index=self.index_name, doc_type=self.index_type, body=item)
            print(res)

    def savadatawww(self,filepath,filename):
        LIST=[]
        global document, documentnode, nodedic
        filedata = xlrd.open_workbook(filepath)
        filetable = filedata.sheet_by_index(0)
        for row in range(0, filetable.nrows):
            for col in range(0, filetable.ncols):
                value = filetable.cell_value(row, col)
                if type(value) == str:
                    value = self.wipe_line_break(value)
                    value = self.delete_BS(value)
                if value == "":
                    continue
                action={
                     "_index": self.index_name,
                     "_type": self.index_type,
                     "_source": {
                     "filename":filename,
                     "filecontene":value.encode('utf-8').decode('utf8')}
                }
                LIST.append(action)
        success, _ = bulk(self.es, LIST, index=self.index_name, raise_on_error=True)
        print('Performed %d actions' % success)

    def savadatadocx(self, filepath, filename):
        LIST = []
        global document, documentnode, nodedic
        file = docx.Document(filepath)
        print(file.paragraphs)
        b = [j.text for j in file.paragraphs]
        value = ''.join(b)
        print(value)
        filenamenew = self.wipe_line_break(filename)
        # for para in file.paragraphs:
        #     print(para.text)
        # for row in range(0, filetable.nrows):
        #     for col in range(0, filetable.ncols):
        #         value = filetable.cell_value(row, col)
        #         if type(value) == str:
        #             value = self.wipe_line_break(value)
        #             value = self.delete_BS(value)
        #         if value == "":
        #             continue
        action = {
            "_index": self.index_name,
            "_type": self.index_type,
            "_source": {
                "filename": filenamenew,
                "filecontene": value.encode('utf-8').decode('utf8')}
        }
        LIST.append(action)
        success, _ = bulk(self.es, LIST, index=self.index_name, raise_on_error=True)
        print('Performed %d actions' % success)
    def get_allfile(self,cwd):
        global row
        global col
        global folderflag
        global subfolder
        get_dir = os.listdir(cwd)
        for i in get_dir:
            print(i)
            sub_dir = os.path.join(cwd, i)
            # print(sub_dir)
            if os.path.isdir(sub_dir):
                subfolder = True
                row += 1
                self.get_allfile(self,sub_dir)
            else:
                result.append(i)
        return result



    def bulk_Index_Data(self):
        '''
        用bulk将批量数据存储到es
        :return:
        '''
        list = [
            {"date": "2017-09-13",
             "source": "慧聪网",
             "link": "http://info.broadcast.hc360.com/2017/09/130859749974.shtml",
             "keyword": "电视",
             "title": "付费 电视 行业面临的转型和挑战"
             },
            {"date": "2017-09-13",
             "source": "中国文明网",
             "link": "http://www.wenming.cn/xj_pd/yw/201709/t20170913_4421323.shtml",
             "keyword": "电视",
             "title": "电视 专题片《巡视利剑》广获好评：铁腕反腐凝聚党心民心"
             },
            {"date": "2017-09-13",
             "source": "人民电视",
             "link": "http://tv.people.com.cn/BIG5/n1/2017/0913/c67816-29533981.html",
             "keyword": "电视",
             "title": "中国第21批赴刚果（金）维和部隊启程--人民 电视 --人民网"
             },
            {"date": "2017-09-13",
             "source": "站长之家",
             "link": "http://www.chinaz.com/news/2017/0913/804263.shtml",
             "keyword": "电视",
             "title": "电视 盒子 哪个牌子好？ 吐血奉献三大选购秘笈"
             }
        ]
        ACTIONS = []
        i = 1
        for line in list:
            action = {
                "_index": self.index_name,
                "_type": self.index_type,
                "_id": i, #_id 也可以默认生成，不赋值
                "_source": {
                    "date": line['date'],
                    "source": line['source'].encode('utf-8').decode('utf8'),
                    "link": line['link'],
                    "keyword": line['keyword'].encode('utf-8').decode('utf8'),
                    "title": line['title'].encode('utf-8').decode('utf8')}
            }
            i += 1
            ACTIONS.append(action)
            # 批量处理
        success, _ = bulk(self.es, ACTIONS, index=self.index_name, raise_on_error=True)
        print('Performed %d actions' % success)

    def Delete_Index_Data(self,id):
        '''
        删除索引中的一条
        :param id:
        :return:
        '''
        doc={"query": {"bool": {"must": [], "must_not": [], "should": []}}, "from": 0, "size": 10000, "sort": [], "aggs": {}}

        _searched = self.es.search(index=self.index_name, doc_type=self.index_type, body=doc)
        # print(_searched)
        res_list=_searched["hits"]["hits"]
        # print(_searched["hits"]["total"])
        for i in range(len(res_list)):
            id = res_list[i]["_id"]
            delete = self.es.delete(self.index_name, doc_type=self.index_type, id=id)
        # res = self.es.delete(index=self.index_name, doc_type=self.index_type, Any)
        # print res

    def Get_Data_Id(self,id):

        res = self.es.get(index=self.index_name, doc_type=self.index_type,id=id)
        print(res['_source'])

        # print '------------------------------------------------------------------'
        #
        # # 输出查询到的结果
        for hit in res['hits']['hits']:
            # print hit['_source']
            print (hit['_source']['date'],hit['_source']['source'],hit['_source']['link'],hit['_source']['keyword'],hit['_source']['title'])

    def Get_Data_By_Body(self):
        # doc = {'query': {'match_all': {}}}
        doc = {
            "query": {
                "match": {
                    "keyword": "电视"
                }
            }
        }
        _searched = self.es.search(index=self.index_name, doc_type=self.index_type, body=doc)

        for hit in _searched['hits']['hits']:
            # print hit['_source']
            print (hit['_source']['date'], hit['_source']['source'], hit['_source']['link'], hit['_source']['keyword'], \
            hit['_source']['title'])




obj =ElasticObj("file","content",ip ="127.0.0.1")
# obj.Delete_Index_Data('*')
# obj = ElasticObj("ott1", "ott_type1")

# obj.create_index()
# obj.bulk_power_data()

# filepathset = obj.get_allfile(path)
# for item in filepathset:
#     # if "油浸式变压器" in item:
#     filepath = path + "\\" + item
#     obj.savadatawww(filepath)
obj.savadatadocx(r'C:\Users\86136\Desktop\电力缺陷\通辽调研资料旧\变电五项管理规定的206册细则Word版\变电检修管理规定细则docx\国家电网公司变电检修管理规定（试行） 第1分册 油浸式变压器（电抗器）检修细则.docx','国家电网公司变电检修管理规定（试行） 第1分册 油浸式变压器（电抗器）检修细则')
# obj.Index_Data()
# obj.bulk_Index_Data()
# obj.IndexData()
# obj.Delete_Index_Data(1)
# csvfile = 'D:/work/ElasticSearch/exportExcels/2017-08-31_info.csv'
# obj.Index_Data_FromCSV(csvfile)
# obj.GetData(es)